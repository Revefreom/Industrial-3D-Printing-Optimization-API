"""
3D Yazıcı Maliyet Hesaplayıcı - Export Modülü
PDF ve Word formatında fatura oluşturma
"""

from datetime import datetime
import os

# PDF için ReportLab
try:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

# Word için python-docx
try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


def format_currency(amount, currency="TRY"):
    """Para birimini formatlar."""
    symbols = {"TRY": "₺", "USD": "$", "EUR": "€"}
    symbol = symbols.get(currency, currency)
    return f"{amount:,.2f} {symbol}"


def export_to_pdf(invoice_data, company_info, output_path):
    """
    Faturayı PDF olarak dışa aktarır.
    
    Args:
        invoice_data: dict - Fatura bilgileri
        company_info: dict - Şirket bilgileri
        output_path: str - Çıktı dosya yolu
    """
    if not REPORTLAB_AVAILABLE:
        raise ImportError("PDF oluşturmak için 'reportlab' kütüphanesi gerekli. "
                         "Yüklemek için: pip install reportlab")
    
    # Türkçe karakter desteği için font kaydet
    font_name = 'Helvetica'  # varsayılan
    font_bold = 'Helvetica-Bold'
    
    try:
        # Windows Arial fontunu kullan (Türkçe desteği var)
        arial_path = "C:/Windows/Fonts/arial.ttf"
        arial_bold_path = "C:/Windows/Fonts/arialbd.ttf"
        
        if os.path.exists(arial_path):
            pdfmetrics.registerFont(TTFont('Arial', arial_path))
            font_name = 'Arial'
        if os.path.exists(arial_bold_path):
            pdfmetrics.registerFont(TTFont('Arial-Bold', arial_bold_path))
            font_bold = 'Arial-Bold'
    except Exception:
        pass  # Varsayılan fontları kullan
    
    doc = SimpleDocTemplate(output_path, pagesize=A4,
                           rightMargin=2*cm, leftMargin=2*cm,
                           topMargin=2*cm, bottomMargin=2*cm)
    
    elements = []
    styles = getSampleStyleSheet()
    
    # Başlık stili
    title_style = ParagraphStyle(
        'TurkishTitle',
        fontName=font_bold,
        fontSize=18,
        spaceAfter=12,
        alignment=1  # Center
    )
    
    normal_style = ParagraphStyle(
        'TurkishNormal',
        fontName=font_name,
        fontSize=10,
        spaceAfter=6
    )
    
    heading2_style = ParagraphStyle(
        'TurkishHeading2',
        fontName=font_bold,
        fontSize=14,
        spaceAfter=10,
        spaceBefore=10
    )
    
    # Şirket Bilgileri
    elements.append(Paragraph(company_info.get('name', 'Sirket Adi'), title_style))
    elements.append(Paragraph(company_info.get('address', ''), normal_style))
    elements.append(Paragraph(f"Tel: {company_info.get('phone', '')} | E-posta: {company_info.get('email', '')}", normal_style))
    elements.append(Paragraph(f"Vergi No: {company_info.get('tax_number', '')}", normal_style))
    elements.append(Spacer(1, 20))
    
    # Fatura Başlığı
    elements.append(Paragraph("FATURA / TEKLIF", title_style))
    elements.append(Spacer(1, 10))
    
    # Fatura ve Müşteri Bilgileri
    invoice_info = [
        ["Fatura No:", invoice_data.get('invoice_number', '-')],
        ["Tarih:", invoice_data.get('date', datetime.now().strftime('%d.%m.%Y'))],
        ["Musteri:", invoice_data.get('customer_name', '-')],
        ["Firma:", invoice_data.get('customer_company', '-')],
        ["Adres:", invoice_data.get('customer_address', '-')],
    ]
    
    info_table = Table(invoice_info, colWidths=[4*cm, 10*cm])
    info_table.setStyle(TableStyle([
        ('FONTNAME', (0, 0), (-1, -1), font_name),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ('FONTNAME', (0, 0), (0, -1), font_bold),
    ]))
    elements.append(info_table)
    elements.append(Spacer(1, 20))
    
    # Baskı Detayları
    elements.append(Paragraph("Baski Detaylari", heading2_style))
    
    print_info = [
        ["Yazici:", invoice_data.get('printer_name', '-')],
        ["Malzeme:", invoice_data.get('material_name', '-')],
        ["Agirlik:", f"{invoice_data.get('print_weight_g', 0)} gram"],
        ["Sure:", f"{invoice_data.get('print_time_minutes', int(invoice_data.get('print_time_hours', 0) * 60))} dakika"],
        ["Baski Tipi:", "Ilk Baski" if invoice_data.get('is_first_print', True) else "Tekrar Baski"],
    ]
    
    print_table = Table(print_info, colWidths=[4*cm, 10*cm])
    print_table.setStyle(TableStyle([
        ('FONTNAME', (0, 0), (-1, -1), font_name),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ('FONTNAME', (0, 0), (0, -1), font_bold),
    ]))
    elements.append(print_table)
    elements.append(Spacer(1, 20))
    
    # Maliyet Dökümü
    elements.append(Paragraph("Maliyet Dokumu", heading2_style))
    
    currency = invoice_data.get('currency', 'TRY')
    cost_data = [
        ["Kalem", "Tutar"],
        ["Malzeme Maliyeti", format_currency(invoice_data.get('material_cost', 0), currency)],
        ["Enerji Maliyeti", format_currency(invoice_data.get('energy_cost', 0), currency)],
        ["Amortisman (Sarf)", format_currency(invoice_data.get('depreciation_cost', 0), currency)],
        ["Hazirlik/Dilimleme", format_currency(invoice_data.get('preparation_cost', 0), currency)],
        ["Hata Riski Payi", format_currency(invoice_data.get('failure_risk_cost', 0), currency)],
    ]
    
    # Özel indirim varsa ekle
    if invoice_data.get('special_discount', 0) != 0:
        discount = invoice_data.get('special_discount', 0)
        discount_note = invoice_data.get('special_discount_note', 'Ozel Indirim')
        cost_data.append([f"{discount_note}", format_currency(-discount, currency)])
    
    cost_data.append(["TOPLAM", format_currency(invoice_data.get('total_cost', 0), currency)])
    
    cost_table = Table(cost_data, colWidths=[10*cm, 4*cm])
    cost_table.setStyle(TableStyle([
        ('FONTNAME', (0, 0), (-1, -1), font_name),
        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (1, 0), (1, -1), 'RIGHT'),
        ('FONTNAME', (0, 0), (-1, 0), font_bold),
        ('FONTNAME', (0, -1), (-1, -1), font_bold),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
        ('TOPPADDING', (0, 0), (-1, -1), 8),
        ('BACKGROUND', (0, -1), (-1, -1), colors.lightgrey),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
    ]))
    elements.append(cost_table)
    
    # Alt bilgi
    elements.append(Spacer(1, 30))
    footer_style = ParagraphStyle('Footer', fontName=font_name, fontSize=8, textColor=colors.grey, alignment=1)
    elements.append(Paragraph("Bu belge 3D Yazici Maliyet Hesaplayici Pro ile olusturulmustur.", footer_style))
    
    doc.build(elements)
    return output_path


def export_to_word(invoice_data, company_info, output_path):
    """
    Faturayı Word olarak dışa aktarır.
    
    Args:
        invoice_data: dict - Fatura bilgileri
        company_info: dict - Şirket bilgileri
        output_path: str - Çıktı dosya yolu
    """
    if not DOCX_AVAILABLE:
        raise ImportError("Word belgesi oluşturmak için 'python-docx' kütüphanesi gerekli. "
                         "Yüklemek için: pip install python-docx")
    
    doc = Document()
    
    # Şirket Başlığı
    title = doc.add_heading(company_info.get('name', 'Şirket Adı'), 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Şirket Bilgileri
    company_para = doc.add_paragraph()
    company_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    company_para.add_run(f"{company_info.get('address', '')}\n").font.size = Pt(10)
    company_para.add_run(f"Tel: {company_info.get('phone', '')} | E-posta: {company_info.get('email', '')}\n").font.size = Pt(10)
    company_para.add_run(f"Vergi No: {company_info.get('tax_number', '')}").font.size = Pt(10)
    
    doc.add_paragraph()
    
    # Fatura Başlığı
    invoice_title = doc.add_heading('FATURA / TEKLİF', level=1)
    invoice_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Fatura Bilgileri Tablosu
    info_table = doc.add_table(rows=5, cols=2)
    info_table.style = 'Table Grid'
    
    info_data = [
        ("Fatura No:", invoice_data.get('invoice_number', '-')),
        ("Tarih:", invoice_data.get('date', datetime.now().strftime('%d.%m.%Y'))),
        ("Müşteri:", invoice_data.get('customer_name', '-')),
        ("Firma:", invoice_data.get('customer_company', '-')),
        ("Adres:", invoice_data.get('customer_address', '-')),
    ]
    
    for i, (label, value) in enumerate(info_data):
        info_table.rows[i].cells[0].text = label
        info_table.rows[i].cells[1].text = str(value)
        info_table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
    
    doc.add_paragraph()
    
    # Baskı Detayları
    doc.add_heading('Baskı Detayları', level=2)
    
    print_table = doc.add_table(rows=5, cols=2)
    print_table.style = 'Table Grid'
    
    print_data = [
        ("Yazıcı:", invoice_data.get('printer_name', '-')),
        ("Malzeme:", invoice_data.get('material_name', '-')),
        ("Ağırlık:", f"{invoice_data.get('print_weight_g', 0)} gram"),
        ("Süre:", f"{invoice_data.get('print_time_minutes', int(invoice_data.get('print_time_hours', 0) * 60))} dakika"),
        ("Baskı Tipi:", "İlk Baskı" if invoice_data.get('is_first_print', True) else "Tekrar Baskı"),
    ]
    
    for i, (label, value) in enumerate(print_data):
        print_table.rows[i].cells[0].text = label
        print_table.rows[i].cells[1].text = str(value)
        print_table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
    
    doc.add_paragraph()
    
    # Maliyet Dökümü
    doc.add_heading('Maliyet Dökümü', level=2)
    
    currency = invoice_data.get('currency', 'TRY')
    
    cost_items = [
        ("Malzeme Maliyeti", invoice_data.get('material_cost', 0)),
        ("Enerji Maliyeti", invoice_data.get('energy_cost', 0)),
        ("Amortisman (Sarf)", invoice_data.get('depreciation_cost', 0)),
        ("Hazırlık/Dilimleme", invoice_data.get('preparation_cost', 0)),
        ("Hata Riski Payı", invoice_data.get('failure_risk_cost', 0)),
    ]
    
    # Özel indirim varsa ekle
    if invoice_data.get('special_discount', 0) != 0:
        discount_note = invoice_data.get('special_discount_note', 'Özel İndirim')
        cost_items.append((discount_note, -invoice_data.get('special_discount', 0)))
    
    cost_items.append(("TOPLAM", invoice_data.get('total_cost', 0)))
    
    cost_table = doc.add_table(rows=len(cost_items) + 1, cols=2)
    cost_table.style = 'Table Grid'
    
    # Başlık satırı
    cost_table.rows[0].cells[0].text = "Kalem"
    cost_table.rows[0].cells[1].text = "Tutar"
    for cell in cost_table.rows[0].cells:
        cell.paragraphs[0].runs[0].bold = True
    
    # Veri satırları
    for i, (label, amount) in enumerate(cost_items):
        cost_table.rows[i + 1].cells[0].text = label
        cost_table.rows[i + 1].cells[1].text = format_currency(amount, currency)
        cost_table.rows[i + 1].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Son satırı (TOPLAM) kalın yap
        if i == len(cost_items) - 1:
            cost_table.rows[i + 1].cells[0].paragraphs[0].runs[0].bold = True
            cost_table.rows[i + 1].cells[1].paragraphs[0].runs[0].bold = True
    
    doc.add_paragraph()
    
    # Alt bilgi
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run("Bu belge 3D Yazıcı Maliyet Hesaplayıcı Pro ile oluşturulmuştur.")
    footer_run.font.size = Pt(8)
    footer_run.font.italic = True
    
    doc.save(output_path)
    return output_path
